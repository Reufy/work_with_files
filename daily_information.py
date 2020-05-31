import textwrap
from docx import Document
from check_text_in_file import check_name_file
from dictionary import *
from copy_file import now_date
from copy_file import copy_file
from parsing import now
import comtypes.client
import comtypes
import re
import os
import shutil


def daily_information():
    arr = []
    ek_book = 0
    date_of = 0
    plot = 0
    decision = 0
    service = 0
    i = 0


    wdFormatDOCX = 16

    copyfile = copy_file(now_date())
    input_dir = copyfile[0]
    output_dir = copyfile[1]
    date_result = copyfile[2]
    print("Создание временных файлов")
    print("----------------------------------------------------------------------")
    for subdir, dirs, files in os.walk(input_dir):
        for file in files:
            if file == "newfile.txt":
                continue
            in_file = os.path.join(subdir, file)
            output_file = file.split('.')[0]
            out_file = output_dir + output_file + ".docx"
            word = comtypes.client.CreateObject('Word.Application')
            # word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatDOCX)
            print("Преобразование файла: " + output_file)
            doc.Close()
            word.Quit()
        break

    # directory_open = "./test/"  # Linux
    directory_open = copyfile[1]  # windows open file
    filename = os.listdir(directory_open)
    f = open(copyfile[0] + 'newfile.txt', 'w+')
    f.write("Файлы за: " + date_result + "\n")
    print("----------------------------------------------------------------------")
    print("Анализ файлов")
    print("----------------------------------------------------------------------")
    for fn in filename:
        row_first = 0
        print("Анализ: " + fn)
        doc_open_name = copyfile[1] + fn  # windows open file
        # doc_open_name = "./test/" + fn  # Linux
        wordDoc = Document(doc_open_name)
        len_par = len(wordDoc.paragraphs)
        temp_name = check_name_file(len_par, f, wordDoc)
        for table in wordDoc.tables:
            for row in table.rows:
                for cell in row.cells:
                    dubtext = cell.text.lower()
                    dubtext = dubtext.replace('\n', ' ').split()
                    dubtext = ' '.join(dubtext)
                    arr.append(dubtext)

                if row_first == 0:
                    for value in arr:
                        if re.findall(r"\bслово", value) or re.findall(r"\bслово.", value) or re.findall(r"\bслово",
                                                                                                   value) or re.findall(
                                r"\bслово", value) or re.findall(r"\bслово", value):
                            if temp_name == 'Слово':
                                ek_book = 0
                            else:
                                ek_book = arr.index(value)
                        elif re.findall(r"\bслово", value.replace(" ", "")) or re.findall(r"слово слово", value.replace(" ", " ")):
                            plot = arr.index(value)
                        elif re.findall(r"\bслово", value):
                            decision = arr.index(value)
                        elif re.findall(r"\bслово", value):
                            service = arr.index(value)
                    row_first = 1
                    arr.clear()
                else:
                    arr[service] = ''
                    for value in arr:
                        for string in other_word:
                            regul = re.findall(string, value)
                            if regul:
                                f.write(str(temp_name))
                                while True:
                                    i += 1
                                    final_text = "№" + str(i) + "  |" + arr[ek_book] + " | " + arr[plot] + " | " \
                                                 + arr[decision] + "| "
                                    if len(final_text) < 110:
                                        f.write("{:<}".format(final_text[:110]))
                                        f.write("\n")
                                    else:
                                        while len(final_text) > 110:
                                            f.write(final_text[:110])
                                            f.write("\n")
                                            final_text = final_text[110:]
                                        f.write(final_text[:110])
                                        f.write("\n")
                                    break
                                break
                    arr.clear()
    text_ = now(now_date())
    f.write("\n" + 'Колличество' + " " + str(i) + '\n')
    for text in text_:
        if len(text) < 110:
            f.write('\n' + "{:<}".format(text[:110]))
        else:
            while len(text) > 110:
                f.write('\n' + text[:110])
                text = text[110:]
            f.write("\n")
            f.write(text[:110])
            f.write("\n")
    f.close()
    del_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), copyfile[1])
    shutil.rmtree(del_path)
    print("----------------------------------------------------------------------")
    print("Удаление временных файлов")
    print("----------------------------------------------------------------------")
    open_result_file = copyfile[0] + 'newfile.txt'
    os.system(open_result_file)

